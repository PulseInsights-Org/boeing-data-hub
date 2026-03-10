"""Shared formatting helpers for .docx generation."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def create_doc(title: str) -> Document:
    """Create a new document with standard styling."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Configure heading styles
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Arial"
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.styles["Heading 1"].font.size = Pt(20)
    doc.styles["Heading 2"].font.size = Pt(16)
    doc.styles["Heading 3"].font.size = Pt(13)

    # Add page numbers in footer
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = parse_xml(
        '<w:fldSimple {} w:instr=" PAGE "/>'.format(nsdecls("w"))
    )
    run._r.append(fld)

    # Title page
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.space_before = Pt(120)
    run = tp.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Boeing Data Hub")
    run.font.size = Pt(16)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("February 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()
    return doc


def add_table(doc: Document, headers: list, rows: list) -> None:
    """Add a bordered table with shaded header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(
            '<w:shd {} w:fill="1A1A2E" w:val="clear"/>'.format(nsdecls("w"))
        )
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            run.font.name = "Arial"
            # Alternate row shading
            if r_idx % 2 == 1:
                shading = parse_xml(
                    '<w:shd {} w:fill="F2F2F2" w:val="clear"/>'.format(nsdecls("w"))
                )
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # spacing


def add_bold_paragraph(doc: Document, text: str) -> None:
    """Add a paragraph with bold text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_tip(doc: Document, text: str) -> None:
    """Add an italic tip paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"Tip: {text}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_screenshot_placeholder(doc: Document, label: str) -> None:
    """Add a screenshot placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[Screenshot: {label}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)


def add_diagram_placeholder(doc: Document, label: str) -> None:
    """Add a diagram placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[Diagram: {label}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)


def add_code_block(doc: Document, code: str) -> None:
    """Add a code block with monospace font and gray background."""
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_numbered_steps(doc: Document, steps: list) -> None:
    """Add numbered steps, bolding text between ** markers."""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        # Parse bold markers
        parts = step.split("**")
        prefix = p.add_run(f"{i}. ")
        prefix.font.name = "Arial"
        prefix.font.size = Pt(11)
        for j, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            if j % 2 == 1:
                run.bold = True


def add_bullet(doc: Document, text: str) -> None:
    """Add a bullet point."""
    doc.add_paragraph(text, style="List Bullet")


def add_file_entry(doc: Document, path: str, desc: str) -> None:
    """Add a file entry with arrow: path -> description."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"\u27a1 {path}")
    run.bold = True
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run2 = p.add_run(f" \u2014 {desc}")
    run2.font.name = "Arial"
    run2.font.size = Pt(10)


def save_doc(doc: Document, filename: str) -> str:
    """Save document and return path."""
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path
