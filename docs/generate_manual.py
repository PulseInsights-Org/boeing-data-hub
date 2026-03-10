"""
Generate Boeing Data Sync User Manual (.docx)
Matches exact structure and styling of the Shopify Doc Portal User Manual.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str):
    """Set cell background shading."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """Add thin borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def styled_table(doc, headers, rows):
    """Create a styled table matching the reference manual."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_table_borders(table)

    # Header row
    for ci, text in enumerate(headers):
        cell = table.rows[0].cells[ci]
        set_cell_shading(cell, "1B4F72")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(10)

    return table


# Abstract numbering counters
_abstract_num_id = [0]


def _get_next_abstract_id():
    _abstract_num_id[0] += 1
    return _abstract_num_id[0]


def _ensure_numbering_part(doc):
    """Ensure the document has a numbering part."""
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        # Create numbering part if it doesn't exist
        from docx.opc.part import Part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        numbering_part = None
    return doc


def add_bullet_list(doc, items):
    """Add a bullet list. Each item can be a string or a tuple (bold_prefix, rest)."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            bold_text, rest_text = item
            run_b = p.add_run(bold_text)
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_r = p.add_run(rest_text)
            run_r.font.size = Pt(10)
        else:
            run = p.add_run(str(item))
            run.font.size = Pt(10)


def add_numbered_list(doc, items):
    """Add a numbered list."""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            bold_text, rest_text = item
            run_b = p.add_run(bold_text)
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_r = p.add_run(rest_text)
            run_r.font.size = Pt(10)
        else:
            run = p.add_run(str(item))
            run.font.size = Pt(10)


def add_normal(doc, text):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_note(doc, text):
    """Add a Note: paragraph (green bold 'Note:' + grey body)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run("Note: ")
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(0x14, 0x8F, 0x77)
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_body.font.size = Pt(10)
    return p


def add_warning(doc, text):
    """Add a Warning: paragraph (red bold 'Warning:' + grey body)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run("Warning: ")
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_body.font.size = Pt(10)
    return p


def add_tip(doc, text):
    """Add a tip paragraph with lightbulb emoji."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\U0001f4a1 " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_screenshot_placeholder(doc, caption):
    """Add a screenshot placeholder."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"[Screenshot: {caption}]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    return p


def set_header_footer(section, header_text, footer_text):
    """Set header and footer for a section."""
    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(header_text)
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(footer_text)
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ---------------------------------------------------------------------------
# Main Document Generation
# ---------------------------------------------------------------------------

def generate_manual():
    doc = Document()

    # -- Page setup (Letter, 1-inch margins) --
    for section in doc.sections:
        section.page_width = Emu(7772400)   # Letter width
        section.page_height = Emu(10058400)  # Letter height
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # -- Style definitions --
    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    h1_style.paragraph_format.space_before = Pt(5)
    h1_style.paragraph_format.space_after = Pt(10)

    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    h2_style.paragraph_format.space_before = Pt(14)
    h2_style.paragraph_format.space_after = Pt(7)

    # Normal
    normal_style = doc.styles['Normal']
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.font.size = Pt(10)

    # ======================================================================
    # COVER PAGE
    # ======================================================================
    # Spacing before title
    for _ in range(8):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BOEING DATA SYNC")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("User Manual")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

    # Spacing
    for _ in range(3):
        doc.add_paragraph()

    # Prepared for
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for Skynet Parts")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Boeing Product Data Synchronization & Publishing Platform")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Spacing
    for _ in range(2):
        doc.add_paragraph()

    # Version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  |  2025")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Page break
    doc.add_page_break()

    # ======================================================================
    # TABLE OF CONTENTS
    # ======================================================================
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    toc_items = [
        "1. Overview",
        "2. System Requirements",
        "3. Getting Started",
        "4. Dashboard Overview",
        "5. Core Features \u2014 Fetch & Process",
        "6. Core Features \u2014 Published Products",
        "7. Core Features \u2014 Auto-Sync",
        "8. Editing Products",
        "9. Logout",
        "10. Troubleshooting",
        "11. FAQ",
        "12. Support Contact",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.size = Pt(10)

    doc.add_page_break()

    # -- Header / Footer for all sections --
    for section in doc.sections:
        set_header_footer(section, "Boeing Data Sync  |  User Manual", "Page 1  |  Confidential")

    # ======================================================================
    # 1. OVERVIEW
    # ======================================================================
    doc.add_heading("1. Overview", level=1)

    doc.add_heading("1.1 Business Context", level=2)
    add_normal(doc,
        "Boeing Data Sync is a centralized product data management and publishing portal "
        "built for Shopify-based aviation parts stores. It connects directly to Boeing\u2019s "
        "Part Number Availability (PNA) API to fetch real-time pricing, inventory, and product "
        "details for aviation parts \u2014 then normalizes and publishes that data directly to "
        "your Shopify storefront."
    )

    doc.add_heading("1.2 Key Capabilities", level=2)
    styled_table(doc,
        ["Capability", "Description"],
        [
            ["Bulk Part Search", "Search Boeing\u2019s catalog for up to 50,000 part numbers in a single request, with automatic batching and real-time progress tracking"],
            ["Data Normalization", "Automatically normalizes Boeing product data including pricing, dimensions, weight, inventory, compliance codes, and images into a retail-ready format"],
            ["Shopify Publishing", "Publish individual or bulk products to your Shopify store with automated price markup, inventory sync, and image upload"],
            ["Published Products Management", "Browse, search, and manage all products published to Shopify with direct admin links"],
            ["Automatic Price & Inventory Sync", "Scheduled hourly synchronization keeps your Shopify product prices and inventory up to date with Boeing\u2019s latest data"],
            ["Real-Time Pipeline Tracking", "Monitor extraction, normalization, and publishing progress in real time with color-coded status badges"],
            ["Product Editing", "Edit product titles, descriptions, and pricing before or after publishing to Shopify"],
        ]
    )

    doc.add_heading("1.3 Intended Audience", level=2)
    add_normal(doc, "This manual is designed for the following users:")
    add_bullet_list(doc, [
        ("Store Administrators", " \u2014 Responsible for managing the product catalog, configuring sync schedules, and overseeing the data pipeline from Boeing to Shopify."),
        ("Staff Members", " \u2014 Day-to-day users who search for Boeing parts, review normalized data, publish products, and monitor sync status."),
    ])

    doc.add_heading("1.4 How to Access", level=2)
    add_normal(doc,
        "Boeing Data Sync is accessed through the Aviation Gateway \u2014 the centralized login "
        "portal for all Skynet Parts applications."
    )
    add_bullet_list(doc, [
        ("Aviation Gateway URL: ", "https://hangar.skynetparts.com/"),
        ("Look for the card labeled: ", "Boeing Data Sync"),
    ])
    add_note(doc, "If you do not have access, contact your system administrator.")

    # ======================================================================
    # 2. SYSTEM REQUIREMENTS
    # ======================================================================
    doc.add_heading("2. System Requirements", level=1)
    add_normal(doc,
        "Before using Boeing Data Sync, ensure your system meets the following requirements "
        "for the best experience."
    )

    doc.add_heading("2.1 Supported Browsers", level=2)
    add_normal(doc,
        "Boeing Data Sync is a web-based application that runs entirely in your browser. "
        "The following browsers have been tested and are officially supported:"
    )
    styled_table(doc,
        ["Browser", "Minimum Version", "Recommended"],
        [
            ["Google Chrome", "90+", "Latest version"],
            ["Mozilla Firefox", "88+", "Latest version"],
            ["Microsoft Edge", "90+", "Latest version"],
            ["Safari", "14+", "Latest version"],
        ]
    )
    add_note(doc, "Internet Explorer is not supported. If you are using Internet Explorer, please switch to one of the supported browsers listed above.")
    add_tip(doc, "For the best performance and compatibility, we recommend keeping your browser updated to the latest version. Automatic updates are enabled by default in most modern browsers.")

    doc.add_heading("2.2 Network Requirements", level=2)
    add_normal(doc,
        "Boeing Data Sync requires network access to communicate with Boeing\u2019s API, "
        "your Shopify store, and the data storage service. Ensure the following conditions are met:"
    )
    add_bullet_list(doc, [
        "A stable internet connection is required to access the portal.",
        "The following domains must be accessible from your network:",
        ("", "https://hangar.skynetparts.com/ \u2014 Aviation Gateway (login)"),
        ("", "Your Boeing Data Sync URL \u2014 provided by your system administrator"),
    ])
    add_normal(doc,
        "If your organization uses a firewall or content filter, ensure these domains are "
        "whitelisted. Blocked network access may result in the portal failing to load, "
        "Boeing searches timing out, or products failing to publish."
    )

    doc.add_heading("2.3 Account Requirements", level=2)
    add_normal(doc,
        "Access to Boeing Data Sync is managed through the Aviation Gateway single sign-on "
        "(SSO) system. To use the portal, you need:"
    )
    add_bullet_list(doc, [
        "An active user account provisioned by your system administrator.",
        "Your account must have access to the Boeing Data Sync application within the Aviation Gateway.",
        "Your login credentials (email and password) for the Aviation Gateway.",
    ])
    add_normal(doc,
        "If your account has not been set up or you have forgotten your password, contact your "
        "system administrator for assistance. Account provisioning and password resets are handled "
        "at the organization level through the Aviation Gateway admin panel."
    )

    doc.add_heading("2.4 Display Recommendations", level=2)
    add_normal(doc, "For the best experience when working with product data and batch operations, the following display settings are recommended:")
    add_bullet_list(doc, [
        ("Screen resolution: ", "1280 x 720 or higher. The dashboard uses a multi-panel layout that benefits from wider screens."),
        ("Zoom level: ", "100% browser zoom is recommended. Higher zoom levels may cause the layout to shift to a stacked view."),
        ("", "The product tables and pipeline tracking panels are optimized for desktop and laptop browsers."),
    ])
    add_note(doc, "Boeing Data Sync is optimized for desktop and laptop browsers. While it may work on tablets, the full dashboard experience is best on screens 13 inches or larger.")

    # ======================================================================
    # 3. GETTING STARTED
    # ======================================================================
    doc.add_heading("3. Getting Started", level=1)
    add_normal(doc,
        "This section walks you through the initial steps to access and begin using Boeing Data Sync, "
        "from logging in to the Aviation Gateway to navigating to the main dashboard."
    )

    doc.add_heading("3.1 Logging In", level=2)
    add_normal(doc,
        "All access to Boeing Data Sync begins at the Aviation Gateway \u2014 the centralized "
        "login portal for all Skynet Parts applications."
    )
    add_numbered_list(doc, [
        "Open your web browser (Chrome, Firefox, or Edge recommended).",
        "Navigate to the Aviation Gateway at: https://hangar.skynetparts.com/",
        "You will see the Welcome Back login screen.",
        "Enter your Email address.",
        "Enter your Password.",
        "Click the Sign In button.",
    ])
    add_screenshot_placeholder(doc, "Aviation Gateway Login Page")
    add_normal(doc,
        "You should now see the Aviation Gateway dashboard with a greeting message and a list "
        "of available portal cards."
    )
    add_tip(doc,
        "Bookmark https://hangar.skynetparts.com/ for quick access. You can also save it "
        "as a shortcut on your desktop or pin it in your browser\u2019s tab bar."
    )
    add_note(doc,
        "If you cannot log in, contact your system administrator to verify your account. "
        "Common login issues include incorrect email address, expired password, or the "
        "account not being provisioned yet."
    )

    doc.add_heading("3.2 Navigating to Boeing Data Sync", level=2)
    add_normal(doc,
        "After logging in to the Aviation Gateway, you will see the dashboard with your "
        "available tools listed as cards. Each card represents a Skynet Parts application "
        "that your account has access to."
    )
    add_numbered_list(doc, [
        "Look for the card labeled Boeing Data Sync.",
        "Click on the Boeing Data Sync card.",
        "You will be redirected to the Boeing Data Sync dashboard.",
    ])
    add_screenshot_placeholder(doc, "Aviation Gateway Dashboard with Boeing Data Sync card highlighted")
    add_normal(doc,
        "You should now see the Boeing Data Sync dashboard with three main tabs: "
        "Fetch & Process, Published Products, and Auto-Sync."
    )
    add_note(doc,
        "If you do not see the Boeing Data Sync card, contact your system administrator "
        "\u2014 you may not have access to this portal. Only users with the appropriate "
        "permissions can see and access the application."
    )

    doc.add_heading("3.3 First-Time Checklist", level=2)
    add_normal(doc, "If this is your first time using Boeing Data Sync, we recommend completing the following steps to familiarize yourself with the system:")
    add_numbered_list(doc, [
        ("Verify your login", " \u2014 Ensure you can log in to the Aviation Gateway and access the Boeing Data Sync card."),
        ("Run a test search", " \u2014 Enter a known Boeing part number in the Fetch & Process tab and click Fetch to verify that Boeing API connectivity is working."),
        ("Review the results", " \u2014 Expand the batch to see the extracted and normalized product data. Verify that pricing, inventory, and dimensions are populated."),
        ("Publish a test product", " \u2014 Select a product with inventory and price, and click Publish to verify that Shopify publishing is working correctly."),
        ("Check Published Products", " \u2014 Switch to the Published Products tab and confirm the product appears with a Shopify link."),
        ("Review Auto-Sync", " \u2014 Open the Auto-Sync tab to see the sync dashboard and verify that the published product has been added to the sync schedule."),
    ])
    add_tip(doc,
        "Completing this checklist ensures that all system components (Boeing API connection, "
        "data normalization, Shopify publishing, and auto-sync) are functioning correctly "
        "before you rely on the system for day-to-day operations."
    )

    # ======================================================================
    # 4. DASHBOARD OVERVIEW
    # ======================================================================
    doc.add_heading("4. Dashboard Overview", level=1)
    add_normal(doc,
        "After navigating to Boeing Data Sync, you will see the main dashboard. This is your "
        "central workspace for managing the Boeing-to-Shopify product pipeline. The dashboard "
        "is organized into a top navigation bar and three main tabs."
    )
    add_screenshot_placeholder(doc, "Main Dashboard")

    doc.add_heading("4.1 How the System Works", level=2)
    add_normal(doc,
        "Boeing Data Sync acts as a bridge between Boeing\u2019s aviation parts catalog and your "
        "Shopify storefront. Understanding the system pipeline will help you use the portal more "
        "effectively."
    )
    add_normal(doc, "The entire flow can be summarized in the following steps:")
    add_numbered_list(doc, [
        "You enter one or more Boeing part numbers into the Fetch & Process tab.",
        "The system queries Boeing\u2019s Part Number Availability (PNA) API to fetch pricing, inventory, dimensions, and product details.",
        "Raw Boeing data is normalized into a retail-friendly format with price markup applied.",
        "You review the normalized products, optionally edit them, and publish to Shopify.",
        "Published products are automatically enrolled in the hourly Auto-Sync schedule.",
        "Every hour, the system re-checks Boeing for price and inventory changes and updates your Shopify store automatically.",
    ])

    doc.add_heading("4.2 Top Navigation Bar", level=2)
    add_normal(doc,
        "The top navigation bar provides quick access to global actions. It is always visible "
        "regardless of which tab you are viewing."
    )
    add_bullet_list(doc, [
        ("Logo", " \u2014 Displays the application name (Boeing Product Normalization & Publishing Dashboard) with an airplane icon."),
        ("Home button", " \u2014 Returns you to the Aviation Gateway dashboard, where you can access other Skynet Parts applications."),
        ("Logout button", " \u2014 Signs you out of the system and ends your session across all connected Skynet Parts applications."),
    ])
    add_note(doc, "The navigation bar is consistent across all tabs. You can always return to the Aviation Gateway or log out from any screen within the portal.")

    doc.add_heading("4.3 Tab Navigation", level=2)
    add_normal(doc, "Below the navigation bar, three tabs organize the main features of the portal:")
    styled_table(doc,
        ["Tab", "Icon", "Description"],
        [
            ["Fetch & Process", "Cloud Download", "Search Boeing for parts, view extraction/normalization progress, and publish products"],
            ["Published Products", "Shopping Bag", "Browse and manage all products that have been published to Shopify"],
            ["Auto-Sync", "Refresh", "Monitor automatic price and inventory synchronization with Boeing"],
        ]
    )
    add_normal(doc,
        "Click on any tab to switch between views. The Published Products tab also shows a "
        "badge with the total number of published products."
    )

    # ======================================================================
    # 5. CORE FEATURES - FETCH & PROCESS
    # ======================================================================
    doc.add_heading("5. Core Features \u2014 Fetch & Process", level=1)

    doc.add_heading("5.1 Entering Part Numbers", level=2)
    add_normal(doc,
        "The Fetch & Process tab is your starting point for bringing Boeing product data into "
        "the system. At the top of the tab, you will see the Fetch Parts input section."
    )
    add_numbered_list(doc, [
        "Click on the input field labeled \u201cEnter part numbers.\u201d",
        "Type or paste one or more Boeing part numbers.",
        "Part numbers can be separated by commas, semicolons, spaces, or new lines.",
        "A counter badge appears showing how many part numbers you have entered (e.g., \u201c3 parts\u201d).",
        "Click the Fetch button to start the search.",
    ])
    add_screenshot_placeholder(doc, "Fetch Parts input with part numbers entered")
    add_tip(doc, "You can paste a large list of part numbers directly from a spreadsheet. The system supports up to 50,000 part numbers in a single request.")
    add_note(doc, "Press Enter to quickly submit your search without clicking the Fetch button (as long as Shift is not held).")

    doc.add_heading("5.2 Understanding the Pipeline", level=2)
    add_normal(doc,
        "When you submit a search, Boeing Data Sync processes the part numbers through a three-stage pipeline:"
    )
    styled_table(doc,
        ["Stage", "Description", "Status Color"],
        [
            ["Extraction", "The system fetches raw product data from Boeing\u2019s PNA API. Part numbers are batched in groups of 10 to respect Boeing\u2019s rate limits.", "Blue"],
            ["Normalization", "Raw Boeing data is parsed and transformed into a retail-ready format: title, description, pricing (with 10% markup), dimensions, weight, inventory, compliance codes, and images.", "Amber/Yellow"],
            ["Publishing", "Normalized products are published to your Shopify store with product images, metafields, and location-based inventory.", "Green"],
        ]
    )

    doc.add_heading("5.3 Monitoring Batch Progress", level=2)
    add_normal(doc,
        "After submitting a search, a batch card appears in the Recent Requests section below "
        "the input field. Each batch card shows real-time progress."
    )
    add_bullet_list(doc, [
        ("Status badge", " \u2014 Shows the current state: pending, processing, completed, failed, or cancelled."),
        ("Progress bar", " \u2014 Visual indicator of how far the batch has progressed."),
        ("Counters", " \u2014 Extracted, Normalized, Published, and Failed counts update in real time."),
        ("Timestamp", " \u2014 When the batch was created."),
        ("Cancel button", " \u2014 Appears for pending or processing batches, allowing you to stop the operation."),
    ])
    add_screenshot_placeholder(doc, "Batch card showing extraction progress")

    doc.add_heading("5.4 Filtering Batches", level=2)
    add_normal(doc, "Use the status filter tabs above the batch list to narrow down which batches are displayed:")
    add_bullet_list(doc, [
        ("All", " \u2014 Shows all batches regardless of status."),
        ("Active", " \u2014 Shows only pending and processing batches."),
        ("Completed", " \u2014 Shows only successfully completed batches."),
        ("Failed", " \u2014 Shows only batches that encountered errors."),
        ("Cancelled", " \u2014 Shows only batches that were manually cancelled."),
    ])

    doc.add_heading("5.5 Expanding Batch Details", level=2)
    add_normal(doc,
        "Click the expand arrow on any batch card to see detailed information about the batch."
    )
    add_normal(doc, "The expanded view includes:")
    add_bullet_list(doc, [
        ("Batch ID", " \u2014 A unique identifier for the batch (useful for support requests)."),
        ("Pipeline Summary Cards", " \u2014 Five cards showing Extracted, Normalized, Published, Failed, and Skipped counts."),
        ("Part Number Tags", " \u2014 Color-coded badges for each part number showing its current status in the pipeline."),
        ("Published Part Numbers", " \u2014 Green tags for parts that have been successfully published to Shopify."),
        ("Skipped Part Numbers", " \u2014 Orange tags for parts that were excluded from publishing (no inventory or no price)."),
        ("Error Details", " \u2014 If any items failed, a table shows the part number and error reason."),
    ])

    doc.add_heading("5.6 Part Number Status Colors", level=2)
    add_normal(doc, "Each part number in a batch is displayed as a color-coded tag indicating its current status:")
    styled_table(doc,
        ["Color", "Status", "Meaning"],
        [
            ["Grey", "Pending", "Waiting to be extracted from Boeing"],
            ["Blue", "Fetched", "Extracted from Boeing, being normalized"],
            ["Amber", "Normalized", "Ready to publish (may have inventory/price issues)"],
            ["Green", "Published", "Successfully published to Shopify"],
            ["Orange", "Blocked", "Skipped \u2014 no inventory, no price, or missing data"],
            ["Red", "Failed", "An error occurred during processing"],
        ]
    )

    doc.add_heading("5.7 Viewing Product Data", level=2)
    add_normal(doc,
        "After normalization completes, you can load and view the product data table for a batch."
    )
    add_numbered_list(doc, [
        "Expand the batch card by clicking the expand arrow.",
        "Click the Load Products button that appears in the action bar.",
        "A product table appears showing all normalized products for this batch.",
    ])
    add_normal(doc, "The product table displays the following columns:")
    add_bullet_list(doc, [
        ("Part Number", " \u2014 The Boeing part number (SKU)"),
        ("Title", " \u2014 Product title derived from the part number"),
        ("Dimensions", " \u2014 Length \u00d7 Width \u00d7 Height with unit of measure"),
        ("Weight", " \u2014 Product weight with unit"),
        ("Price", " \u2014 Cost price from Boeing (before markup)"),
        ("Inventory", " \u2014 Available quantity from Boeing"),
        ("Status", " \u2014 Current pipeline status (fetched, normalized, published, blocked)"),
        ("Actions", " \u2014 Edit, Publish, and Expand buttons for each row"),
    ])
    add_tip(doc, "You can toggle the product table visibility using the Show Table / Hide Table button without losing the loaded data. This is useful for keeping the batch view compact while maintaining real-time status updates.")

    doc.add_heading("5.8 Publishing a Single Product", level=2)
    add_normal(doc, "To publish an individual product to Shopify:")
    add_numbered_list(doc, [
        "Load the product table for a batch (see Section 5.7).",
        "Find the product you want to publish in the table.",
        "Click the Publish button in the Actions column for that product.",
        "A spinner will appear while the product is being published.",
        "Once complete, the status badge changes to \u201cpublished\u201d (green).",
    ])
    add_note(doc, "Products must have both inventory (> 0) and a price (> 0) to be eligible for publishing. Products without these will show a \u201cblocked\u201d status.")

    doc.add_heading("5.9 Bulk Publishing All Products", level=2)
    add_normal(doc,
        "After a batch has been normalized, you can publish all eligible products at once."
    )
    add_numbered_list(doc, [
        "Load the product table for a completed normalization batch.",
        "Click the Publish All button that appears in the action bar.",
        "The button shows the count of publishable products (e.g., \u201cPublish All (12)\u201d).",
        "The system queues all eligible products for publishing.",
        "Progress is tracked in real time \u2014 part number tags turn green as each product is published.",
    ])
    add_normal(doc, "The system automatically filters out products that are:")
    add_bullet_list(doc, [
        "Already published to Shopify",
        "Blocked (no inventory or no price)",
        "Already in a failed state",
    ])
    add_tip(doc, "Products that are skipped during bulk publishing will appear with orange \u201cSkipped\u201d tags in the expanded batch details.")

    doc.add_heading("5.10 Cancelling a Batch", level=2)
    add_normal(doc, "You can cancel an in-progress batch to stop further processing.")
    add_numbered_list(doc, [
        "Find the batch card in the Recent Requests list.",
        "Click the Cancel button that appears on the right side of the batch card.",
        "The batch status changes to \u201ccancelled.\u201d",
        "Products that were already processed remain in the system; only pending items are skipped.",
    ])
    add_warning(doc, "Cancelling a batch cannot be undone. Already-extracted and published products are not affected, but pending items will not be processed.")

    # ======================================================================
    # 6. CORE FEATURES - PUBLISHED PRODUCTS
    # ======================================================================
    doc.add_heading("6. Core Features \u2014 Published Products", level=1)

    doc.add_heading("6.1 Browsing Published Products", level=2)
    add_normal(doc,
        "The Published Products tab shows all products that have been successfully published "
        "to your Shopify store. Products are displayed in a searchable, paginated table."
    )
    add_screenshot_placeholder(doc, "Published Products tab with product table")

    doc.add_heading("6.2 Searching Published Products", level=2)
    add_normal(doc,
        "The search bar at the top of the Published Products tab lets you quickly find a "
        "specific product."
    )
    add_numbered_list(doc, [
        "Click on the search bar at the top of the tab.",
        "Type a part number to search.",
        "The list filters in real time as you type.",
        "To clear the search, delete the text in the search bar.",
    ])

    doc.add_heading("6.3 Product Table Columns", level=2)
    add_normal(doc, "The published products table displays the following information for each product:")
    styled_table(doc,
        ["Column", "Description"],
        [
            ["Part Number", "The Boeing SKU (displayed in monospace font)"],
            ["Title", "Product title as it appears in Shopify"],
            ["Vendor", "The supplier name (typically Boeing Distribution)"],
            ["Price", "The Shopify selling price (with markup applied)"],
            ["Inventory", "Current inventory quantity (green if > 0)"],
            ["Shopify", "A \u201cView\u201d link that opens the product in Shopify admin"],
            ["Updated", "The date and time the product was last updated"],
        ]
    )

    doc.add_heading("6.4 Viewing Product Details", level=2)
    add_normal(doc, "Click the expand arrow on any product row to see additional details:")
    add_bullet_list(doc, [
        ("Cost per Item", " \u2014 The original Boeing cost before markup"),
        ("Dimensions", " \u2014 Length \u00d7 Width \u00d7 Height with unit of measure"),
        ("Weight", " \u2014 Product weight with unit"),
        ("Country of Origin", " \u2014 Where the part was manufactured"),
        ("Description", " \u2014 The full product description as stored in Shopify"),
        ("Product Image", " \u2014 The product image (if available)"),
    ])

    doc.add_heading("6.5 Viewing in Shopify Admin", level=2)
    add_normal(doc,
        "Each published product has a \u201cView\u201d link in the Shopify column. Clicking this "
        "link opens the product directly in your Shopify admin panel in a new browser tab."
    )
    add_note(doc, "You must be logged in to your Shopify admin to access the product page. If you are not logged in, Shopify will prompt you to sign in first.")

    doc.add_heading("6.6 Loading More Products", level=2)
    add_normal(doc,
        "Products are loaded 50 at a time. If you have more than 50 published products, "
        "a \u201cLoad More\u201d button appears at the bottom of the table showing how many "
        "products are loaded out of the total (e.g., \u201cLoad More (50 of 248)\u201d)."
    )

    # ======================================================================
    # 7. CORE FEATURES - AUTO-SYNC
    # ======================================================================
    doc.add_heading("7. Core Features \u2014 Auto-Sync", level=1)
    add_normal(doc,
        "The Auto-Sync tab provides a comprehensive dashboard for monitoring and managing "
        "automatic product synchronization between Boeing and your Shopify store."
    )

    doc.add_heading("7.1 How Auto-Sync Works", level=2)
    add_normal(doc,
        "When a product is published to Shopify through Boeing Data Sync, it is automatically "
        "enrolled in the hourly sync schedule. The system then keeps it up to date without "
        "any manual intervention."
    )
    add_normal(doc, "The sync process works as follows:")
    add_numbered_list(doc, [
        "Published products are distributed across 24 hourly time slots (one per hour of the day).",
        "Every hour, the system fetches the latest pricing and inventory data from Boeing for the products assigned to that hour.",
        "The system compares the new data with the previous data using a hash-based change detection algorithm.",
        "If prices or inventory have changed, the system automatically updates your Shopify store.",
        "If no changes are detected, the product is marked as synced with no Shopify update needed.",
        "Failed syncs are automatically retried every 4 hours.",
    ])
    add_tip(doc, "The sync runs automatically in the background. You do not need to keep the dashboard open for syncing to occur.")

    doc.add_heading("7.2 Status Cards", level=2)
    add_normal(doc,
        "At the top of the Auto-Sync dashboard, six status cards provide an at-a-glance "
        "summary of your sync operations:"
    )
    styled_table(doc,
        ["Card", "Description"],
        [
            ["Total Products", "The total number of products enrolled in the sync schedule"],
            ["Active / Inactive", "How many products are actively syncing vs. deactivated due to repeated failures"],
            ["Success Rate", "The percentage of successful sync operations (color-coded: green \u2265 90%, amber \u2265 70%, red < 70%)"],
            ["Current Hour", "Which hourly time slot is currently being processed"],
            ["Syncing Now / Pending", "Products actively syncing right now and products waiting in queue"],
            ["Failures", "Total number of products with sync failures (highlighted in red if > 0)"],
        ]
    )

    doc.add_heading("7.3 Overview Sub-Tab", level=2)
    add_normal(doc, "The Overview sub-tab (default view) provides three sections:")

    p = doc.add_paragraph()
    run = p.add_run("Hourly Distribution Chart")
    run.bold = True
    run.font.size = Pt(10)
    add_normal(doc,
        "A bar chart showing how your products are distributed across the 24 hourly time slots. "
        "The current hour is highlighted. This helps you visualize load distribution and identify "
        "empty or overloaded time slots."
    )

    p = doc.add_paragraph()
    run = p.add_run("Recent Activity")
    run.bold = True
    run.font.size = Pt(10)
    add_normal(doc,
        "A preview of the last 5 sync operations showing status badges (success/failed), "
        "SKU, price, and quantity for each."
    )

    p = doc.add_paragraph()
    run = p.add_run("Sync Configuration")
    run.bold = True
    run.font.size = Pt(10)
    add_normal(doc,
        "Displays the current sync settings: Sync Mode (Testing or Production), Total Buckets, "
        "Boeing Rate Limit (2 requests/minute), and Batch Size (10 SKUs per API call)."
    )

    doc.add_heading("7.4 Sync History Sub-Tab", level=2)
    add_normal(doc,
        "The Sync History sub-tab shows a detailed table of all recent sync operations. "
        "Each row includes the SKU, sync status (success or failed), timestamp, last synced "
        "price, and last synced quantity."
    )
    add_screenshot_placeholder(doc, "Sync History table")

    doc.add_heading("7.5 Failures Sub-Tab", level=2)
    add_normal(doc,
        "The Failures sub-tab lists all products that have experienced sync failures."
    )
    add_normal(doc, "For each failed product, you can see:")
    add_bullet_list(doc, [
        ("Failure count", " \u2014 How many consecutive failures have occurred (e.g., \u201c3x\u201d)."),
        ("SKU", " \u2014 The product\u2019s part number."),
        ("Deactivated badge", " \u2014 If the product has been automatically deactivated after 5 consecutive failures."),
        ("Reactivate button", " \u2014 Manually reactivate a deactivated product (resets the failure counter to zero)."),
        ("Trigger Sync button", " \u2014 Force an immediate sync for this product, bypassing the hourly schedule."),
    ])
    add_note(doc,
        "Products are automatically deactivated after 5 consecutive sync failures to prevent "
        "repeated errors. You can reactivate them at any time using the Reactivate button."
    )

    doc.add_heading("7.6 Auto-Refresh", level=2)
    add_normal(doc,
        "The Auto-Sync dashboard automatically refreshes every 30 seconds to show the latest "
        "sync status. You can see the auto-refresh indicator at the bottom of the page "
        "(a pulsing green dot when active)."
    )
    add_bullet_list(doc, [
        "Click the Pause button (top right) to stop auto-refresh.",
        "Click the Play button to resume auto-refresh.",
        "Click the Refresh button to manually refresh the dashboard at any time.",
    ])

    # ======================================================================
    # 8. EDITING PRODUCTS
    # ======================================================================
    doc.add_heading("8. Editing Products", level=1)

    doc.add_heading("8.1 Opening the Edit Modal", level=2)
    add_normal(doc, "You can edit a product\u2019s details before or after publishing it to Shopify.")
    add_numbered_list(doc, [
        "Load the product table for a batch in the Fetch & Process tab.",
        "Find the product you want to edit.",
        "Click the Edit button in the Actions column.",
        "The Edit Product modal opens with the current product data.",
    ])
    add_screenshot_placeholder(doc, "Edit Product Modal")

    doc.add_heading("8.2 Editable Fields", level=2)
    add_normal(doc, "The Edit Product modal allows you to modify the following fields:")
    styled_table(doc,
        ["Section", "Field", "Description"],
        [
            ["Basic Information", "Product Title", "The product name as it will appear in Shopify"],
            ["Basic Information", "Description", "The full product description (supports multi-line text)"],
            ["Basic Information", "Manufacturer", "The manufacturer or supplier name"],
            ["Dimensions", "Length / Width / Height", "Physical dimensions of the product"],
            ["Dimensions", "Weight", "Product weight"],
            ["Pricing & Inventory", "Price (USD)", "The selling price"],
            ["Pricing & Inventory", "Inventory Count", "The available quantity"],
        ]
    )

    doc.add_heading("8.3 Saving Changes", level=2)
    add_numbered_list(doc, [
        "Make your changes in the edit form.",
        "Click the Save Changes button.",
        "A loading spinner appears while the changes are being saved.",
        "The modal closes and the product table updates with the new values.",
    ])
    add_note(doc, "If the product has already been published to Shopify, the changes are immediately pushed to your Shopify store as well.")

    # ======================================================================
    # 9. LOGOUT
    # ======================================================================
    doc.add_heading("9. Logout", level=1)
    add_normal(doc,
        "When you are finished working in Boeing Data Sync, it is good practice to log out "
        "of the system to protect your account and data."
    )

    doc.add_heading("9.1 How to Log Out", level=2)
    add_numbered_list(doc, [
        "Click the Logout button (power icon) in the top navigation bar.",
        "Your session will be ended and you will be redirected to the Aviation Gateway login page.",
    ])
    add_normal(doc, "You should see the Aviation Gateway login screen after logging out.")

    doc.add_heading("9.2 Shared Session Behavior", level=2)
    add_normal(doc,
        "Boeing Data Sync uses a shared authentication session with all other Skynet Parts "
        "applications connected through the Aviation Gateway. This means:"
    )
    add_bullet_list(doc, [
        "Logging out of Boeing Data Sync will also sign you out of all other connected Skynet Parts applications (such as Shopify Doc Portal or RFQ Automation Portal).",
        "Logging out from another Skynet Parts application will also end your Boeing Data Sync session.",
        ("Session timeout", " \u2014 If your session expires due to inactivity, you will be redirected to the login page the next time you interact with the portal."),
    ])
    add_note(doc,
        "This shared session behavior is by design. All Skynet Parts applications share a "
        "single sign-on (SSO) session through the Aviation Gateway. If you are unexpectedly "
        "logged out, simply navigate to https://hangar.skynetparts.com/ and sign in again."
    )

    doc.add_heading("9.3 Security Best Practices", level=2)
    add_normal(doc, "To keep your account and product data secure, follow these recommendations:")
    add_bullet_list(doc, [
        "Always log out when using a shared or public computer.",
        "Do not share your Aviation Gateway login credentials with others.",
        "If you suspect unauthorized access to your account, contact your system administrator immediately to reset your password.",
        "Close all browser tabs related to Skynet Parts applications after logging out to fully clear your session.",
    ])
    add_tip(doc,
        "If you frequently switch between multiple Skynet Parts applications, you do not "
        "need to log in separately for each one. A single login at the Aviation Gateway grants "
        "access to all applications your account is authorized for."
    )

    # ======================================================================
    # 10. TROUBLESHOOTING
    # ======================================================================
    doc.add_heading("10. Troubleshooting", level=1)
    add_normal(doc,
        "This section covers common issues and their solutions. If the steps below do not "
        "resolve your problem, contact your system administrator."
    )

    doc.add_heading("10.1 Common Issues & Solutions", level=2)
    styled_table(doc,
        ["Problem", "Possible Cause", "Solution"],
        [
            [
                "No results after searching for part numbers",
                "Part numbers are invalid or not found in Boeing\u2019s catalog",
                "Verify the part numbers are correct. Check for typos or extra characters. Boeing\u2019s catalog may not have data for all part numbers."
            ],
            [
                "Boeing Data Sync card missing from Aviation Gateway",
                "Your account does not have access to this portal",
                "Contact your system administrator to request access to the Boeing Data Sync application."
            ],
            [
                "Batch stuck in \u201cprocessing\u201d state",
                "Boeing API rate limiting or temporary connectivity issue",
                "Wait a few minutes for the system to retry. If it persists for more than 30 minutes, cancel the batch and try again."
            ],
            [
                "Product shows \u201cblocked\u201d status",
                "The product has no inventory (quantity = 0) or no price from Boeing",
                "Blocked products cannot be published. This is expected for out-of-stock or discontinued parts. Check the product details for missing data."
            ],
            [
                "Cannot publish a product",
                "Product is missing price or inventory, or already published",
                "Ensure the product has both a price > 0 and inventory > 0. If it\u2019s already published, check the Published Products tab."
            ],
            [
                "Published product not appearing in Shopify",
                "Publishing is still in progress or failed",
                "Check the batch status for errors. Click the expand arrow to see failed items and error details."
            ],
            [
                "Auto-Sync shows \u201cDeactivated\u201d product",
                "Product failed to sync 5 consecutive times",
                "Go to Auto-Sync > Failures tab and click Reactivate to re-enable syncing. Use Trigger Sync to force an immediate sync."
            ],
            [
                "Logged out unexpectedly",
                "Session expired, or logout from another Skynet Parts app",
                "Navigate to https://hangar.skynetparts.com/ and sign in again."
            ],
            [
                "Prices in Shopify don\u2019t match Boeing",
                "A 10% markup is automatically applied",
                "Boeing Data Sync applies a standard markup to all prices before publishing to Shopify. The original Boeing price is shown as \u201cCost per Item\u201d in the product details."
            ],
        ]
    )

    doc.add_heading("10.2 General Troubleshooting Steps", level=2)
    add_normal(doc,
        "If you encounter an issue not listed above, try the following general steps before "
        "contacting support:"
    )
    add_numbered_list(doc, [
        ("Refresh the page", " \u2014 Press F5 or click the refresh button in your browser. This resolves most temporary display issues."),
        ("Clear your browser cache", " \u2014 Cached data can sometimes cause the portal to display outdated information or fail to load properly."),
        ("Try a different browser", " \u2014 If the issue persists, try accessing the portal in a different supported browser (Chrome, Firefox, or Edge)."),
        ("Check your internet connection", " \u2014 Ensure you have a stable connection. The portal requires network access to communicate with Boeing and Shopify."),
        ("Log out and log back in", " \u2014 Ending your session and starting a new one can resolve authentication and session-related issues."),
        ("Contact your system administrator", " \u2014 If none of the above steps resolve the issue, provide your administrator with the following details: the steps that led to the issue, any error messages displayed, the browser and version you are using, and the approximate time the issue occurred."),
    ])
    add_tip(doc,
        "When reporting an issue, screenshots are extremely helpful. Most browsers allow you "
        "to take a screenshot by pressing the Print Screen key or using the built-in screenshot tool."
    )

    # ======================================================================
    # 11. FAQ
    # ======================================================================
    doc.add_heading("11. FAQ", level=1)
    add_normal(doc,
        "This section answers the most frequently asked questions about Boeing Data Sync. "
        "If your question is not covered here, refer to the relevant section of this manual "
        "or contact your system administrator."
    )

    faqs = [
        (
            "Q1: How many part numbers can I search at once?",
            "Boeing Data Sync supports up to 50,000 part numbers in a single search request. "
            "The system automatically splits them into batches of 10 for the Boeing API and "
            "processes them asynchronously. You can monitor progress in real time through the "
            "batch cards in the Fetch & Process tab."
        ),
        (
            "Q2: Why is a 10% markup applied to prices?",
            "Boeing Data Sync automatically applies a standard markup factor (currently 10%) "
            "to all Boeing prices before publishing to Shopify. The original Boeing price is "
            "preserved as the \u201cCost per Item\u201d in Shopify, while the selling price "
            "reflects the markup. This markup is configured at the system level and cannot be "
            "changed through the portal interface."
        ),
        (
            "Q3: What does \u201cblocked\u201d status mean?",
            "A product is marked as \u201cblocked\u201d when it has no inventory (quantity = 0), "
            "no price, or is missing critical data from Boeing. Blocked products cannot be "
            "published to Shopify. This is expected for out-of-stock or discontinued parts."
        ),
        (
            "Q4: How often does Auto-Sync update prices and inventory?",
            "In production mode, products are synced once every 24 hours. Products are distributed "
            "across 24 hourly time slots, so each product is synced once per day. The system only "
            "updates Shopify when actual price or inventory changes are detected, minimizing "
            "unnecessary API calls."
        ),
        (
            "Q5: What happens if a product fails to sync?",
            "The system automatically retries failed syncs every 4 hours. If a product fails 5 "
            "consecutive times, it is automatically deactivated to prevent repeated errors. "
            "You can manually reactivate it from the Auto-Sync > Failures tab and trigger an "
            "immediate sync."
        ),
        (
            "Q6: Can I edit a product after it\u2019s been published?",
            "Yes. You can edit product titles, descriptions, and pricing through the Edit "
            "Product modal. Changes to already-published products are immediately pushed to "
            "your Shopify store."
        ),
        (
            "Q7: Why does a part number show a variant suffix (e.g., =K3)?",
            "Boeing part numbers sometimes include variant suffixes (like \u201cWF338109=K3\u201d). "
            "The system stores the full variant part number internally for accurate Boeing API "
            "lookups, but strips the suffix when publishing to Shopify (e.g., the SKU in Shopify "
            "becomes \u201cWF338109\u201d)."
        ),
        (
            "Q8: Can multiple staff members use the portal at the same time?",
            "Yes. Multiple users can be logged in and using Boeing Data Sync simultaneously. "
            "Each user accesses the same set of products and batches. Changes made by one user "
            "(such as publishing or editing) are visible to all other users through real-time "
            "updates."
        ),
    ]

    for question, answer in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(question)
        run.bold = True
        run.font.size = Pt(10)
        add_normal(doc, answer)

    # ======================================================================
    # 12. SUPPORT CONTACT
    # ======================================================================
    doc.add_heading("12. Support Contact", level=1)
    add_normal(doc,
        "If you need help with Boeing Data Sync, the following resources and contacts are "
        "available to assist you."
    )

    doc.add_heading("12.1 Getting Help", level=2)
    add_normal(doc, "For technical support, account issues, or questions about the portal, use the following channels:")
    add_bullet_list(doc, [
        ("Email: ", "[TODO: Verify support email]"),
        ("Internal: ", "Contact your system administrator for account provisioning, access requests, password resets, and API connection issues."),
    ])

    doc.add_heading("12.2 What to Include in a Support Request", level=2)
    add_normal(doc, "To help resolve your issue as quickly as possible, include the following information when contacting support:")
    add_bullet_list(doc, [
        "Your name and email address",
        "A description of the issue, including the steps that led to it",
        "Any error messages displayed on screen",
        "The batch ID (if applicable \u2014 found in the expanded batch details)",
        "The browser and version you are using (e.g., Chrome 120)",
        "Screenshots or screen recordings, if available",
        "The approximate date and time the issue occurred",
    ])
    add_tip(doc, "The more detail you provide, the faster the support team can diagnose and resolve your issue. Screenshots are especially helpful for UI-related problems.")

    doc.add_heading("12.3 Additional Resources", level=2)
    add_normal(doc, "Beyond this manual, the following resources may be helpful:")
    add_bullet_list(doc, [
        ("Aviation Gateway: ", "https://hangar.skynetparts.com/ \u2014 The centralized login and application hub for all Skynet Parts tools."),
        ("Shopify Admin: ", "Use the Shopify admin dashboard to manage orders, products, customers, and inventory that are synced from Boeing Data Sync."),
        ("This User Manual: ", "Refer back to the relevant sections of this document for step-by-step guidance on any feature."),
    ])

    # End of Document
    add_normal(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("End of Document")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # ======================================================================
    # Save
    # ======================================================================
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "User_Manual.docx")
    doc.save(output_path)
    print(f"User Manual saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_manual()
